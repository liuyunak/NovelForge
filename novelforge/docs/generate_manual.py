#!/usr/bin/env python3
"""
NovelForge 用户操作手册生成器
生成 Word 格式的完整操作手册（含部署脚本）
"""

import os
import sys

def main():
    try:
        import docx
    except ImportError:
        print("正在安装 docx 库...")
        os.system(f"{sys.executable} -m pip install docx -q")
        import docx
    
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    def add_code_block(doc, text):
        """Add a code snippet with monospace font and gray background"""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        # Add gray border/background
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement('w:bdr')
        fill = OxmlElement('w:fill')
        fill.set(qn('w:val'), 'F2F2F2')
        bdr.append(fill)
        pPr.append(bdr)
        
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        return p
    
    doc = Document()
    
    # Set global styles
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    
    # Set heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = '微软雅黑'
        heading_style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        if i == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.bold = True
        elif i == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
        else:
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True
    
    # ==================== COVER PAGE ====================
    for _ in range(4):
        doc.add_paragraph('')
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('NovelForge')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AI 辅助长篇网文创作工作台')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    
    doc.add_paragraph('')
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('v3.5 | 完整操作手册')
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph('')
    
    detail = doc.add_paragraph()
    detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = detail.add_run('从零开发到部署使用 | 小白友好版\n附带一键部署脚本')
    run.font.size = Pt(12)
    run.font.italic = True
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('2026 年 6 月')
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    doc.add_heading('目录', level=1)
    toc_items = [
        ('第一章', '项目介绍与核心能力'),
        ('第二章', '环境准备与安装部署'),
        ('第三章', 'AI 供应商配置（核心功能）'),
        ('第四章', '创建作品与风格注入'),
        ('第五章', '日常创作流程'),
        ('第六章', '本地模型部署（进阶）'),
        ('第七章', '微调训练（可选）'),
        ('第八章', '差异化功能'),
        ('第九章', '常见问题与故障排查'),
        ('第十章', '附录：一键部署脚本'),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num} {title_text}')
        run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ==================== CHAPTER 1 ====================
    doc.add_heading('第一章 项目介绍与核心能力', level=1)
    
    doc.add_heading('1.1 什么是 NovelForge？', level=2)
    doc.add_paragraph(
        'NovelForge 是一个基于 AI 的网文创作辅助工具，专为 10 万字以上的长篇创作设计。'
        '它不是替代你写作，而是帮你执行——你负责创意和决策，AI 负责生成草稿和检查质量。'
    )
    
    doc.add_heading('1.2 核心能力', level=2)
    core_features = [
        ('全文记忆', 'DeepSeek V4 1M 上下文，50 章+一致性，彻底解决衰减问题'),
        ('风格迁移', '从你的旧稿中提取风格指纹，让 AI 学习你的写作风格'),
        ('DAG 并行编排', '10+ 个 AI Agent 协作，总耗时 35-60 秒生成一章'),
        ('多供应商支持', 'OpenAI、DeepSeek、Ollama、llama.cpp、LM Studio'),
        ('本地微调', '支持 RTX 3070 级别 GPU 进行 LoRA 微调'),
        ('一键封面生成', '根据小说内容自动生成封面图'),
        ('短剧导出', '将章节转换为短剧/漫剧剧本格式'),
    ]
    for title_text, desc in core_features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title_text + ': ').bold = True
        p.add_run(desc)
    
    doc.add_heading('1.3 目标用户', level=2)
    users = [
        '网文作者（日更 4000-6000 字）',
        '10 万字以上长篇创作',
        '有一台普通电脑即可（可选 GPU 用于本地微调）',
        '愿意为创作质量投入少量 API 费用（$2-4/月）',
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')
    
    doc.add_heading('1.4 人机协作创作流程', level=2)
    doc.add_paragraph(
        'NovelForge 采用"AI 执行，人工决策"的设计哲学：\n\n'
        '1. 作者创建新书 + 风格注入 → 系统自动生成设定\n'
        '2. AI 生成章纲 → 作者审批\n'
        '3. AI 流式生成正文 → 作者实时查看\n'
        '4. AI 快速检查 + 深度审计 → 作者决定如何处理\n'
        '5. AI 提取事实 + 润色 → 作者最终审阅\n'
        '6. 作者批准发布 → 系统更新全文记忆\n\n'
        '整个流程有 3 个关键决策点，作者始终掌握控制权。'
    )
    
    doc.add_page_break()
    
    # ==================== CHAPTER 2 ====================
    doc.add_heading('第二章 环境准备与安装部署', level=1)
    
    doc.add_heading('2.1 系统要求', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('配置项', '最低要求'),
        ('操作系统', 'Windows 10+, macOS 12+, Ubuntu 20+'),
        ('内存', '8 GB（推荐 16 GB）'),
        ('磁盘空间', '5 GB 可用空间'),
        ('Node.js', 'v20.0 或更高版本'),
        ('GPU（可选）', 'RTX 3070 8GB（用于本地微调）'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_heading('2.2 安装 Node.js', level=2)
    doc.add_paragraph('步骤 1：下载 Node.js')
    doc.add_paragraph('访问 https://nodejs.org', style='List Bullet')
    doc.add_paragraph('下载 LTS 版本（推荐 v20.x）', style='List Bullet')
    doc.add_paragraph('运行安装程序，一直点"下一步"即可', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('步骤 2：验证安装')
    doc.add_paragraph('在命令行（终端）运行：')
    add_code_block(doc, 'node --versions')
    doc.add_paragraph('应该显示 v20.x.x')
    
    doc.add_heading('2.3 安装 pnpm', level=2)
    doc.add_paragraph('在命令行运行：')
    add_code_block(doc, 'npm install -g pnpm')
    
    doc.add_heading('2.4 安装 Git（可选）', level=2)
    doc.add_paragraph('访问 https://git-scm.com 下载安装', style='List Bullet')
    add_code_block(doc, 'git --version')
    
    doc.add_heading('2.5 下载/克隆项目', level=2)
    doc.add_paragraph('方式 A：使用 Git（推荐）')
    add_code_block(doc, 'git clone https://github.com/your-repo/novelforge.git')
    
    doc.add_paragraph('')
    doc.add_paragraph('方式 B：直接下载 ZIP')
    doc.add_paragraph('访问项目仓库页面', style='List Bullet')
    doc.add_paragraph('点击 Code 按钮 → Download ZIP', style='List Bullet')
    doc.add_paragraph('解压到你想安装的目录', style='List Bullet')
    
    doc.add_heading('2.6 安装依赖', level=2)
    doc.add_paragraph('打开命令行，进入项目目录：')
    add_code_block(doc, 'cd novelforge')
    add_code_block(doc, 'pnpm install')
    doc.add_paragraph('等待安装完成（约 5-10 分钟）')
    
    doc.add_heading('2.7 配置环境变量', level=2)
    doc.add_paragraph('步骤 1：复制环境变量模板')
    add_code_block(doc, 'cp .env.example .env')
    doc.add_paragraph('Windows 用户请手动复制 .env.example 文件并重命名为 .env')
    
    doc.add_paragraph('')
    doc.add_paragraph('步骤 2：编辑 .env 文件')
    doc.add_paragraph('用记事本或 VS Code 打开 .env 文件，修改以下内容：')
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('配置项', '说明'),
        ('DEEPSEEK_API_KEY', '必填！在 https://platform.deepseek.com 注册并获取 API Key'),
        ('DEEPSEEK_BASE_URL', '默认 https://api.deepseek.com，一般不用改'),
        ('LOCAL_MODEL_ENABLED', '如果要用本地模型设为 true，否则 false'),
        ('PORT', '后端服务端口，默认 3001'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph('')
    doc.add_paragraph('示例 .env 文件内容：')
    add_code_block(doc, '# 必填：DeepSeek API Key')
    add_code_block(doc, 'DEEPSEEK_API_KEY=sk-your-api-key-here')
    add_code_block(doc, 'DEEPSEEK_BASE_URL=https://api.deepseek.com')
    
    doc.add_heading('2.8 初始化项目', level=2)
    add_code_block(doc, 'pnpm run init')
    doc.add_paragraph('这会自动创建必要的目录结构')
    
    doc.add_heading('2.9 启动服务', level=2)
    doc.add_paragraph('方式 A：使用一键启动脚本（推荐）')
    add_code_block(doc, '双击 start.bat')
    
    doc.add_paragraph('')
    doc.add_paragraph('方式 B：使用 pnpm 命令')
    add_code_block(doc, 'pnpm dev')
    
    doc.add_paragraph('')
    doc.add_paragraph('启动成功后，浏览器会自动打开 http://localhost:3000')
    doc.add_paragraph('后端 API 运行在 http://localhost:3001')
    
    doc.add_paragraph('')
    doc.add_heading('2.10 停止服务', level=2)
    add_code_block(doc, '双击 stop.bat')
    
    doc.add_page_break()
    
    # ==================== CHAPTER 3 ====================
    doc.add_heading('第三章 AI 供应商配置（核心功能）', level=1)
    
    doc.add_heading('3.1 什么是 AI 供应商？', level=2)
    doc.add_paragraph(
        'NovelForge 支持多种 AI 后端，称为"供应商"。你可以同时配置多个供应商，'
        '并为不同的写作 Agent 选择不同的供应商。'
    )
    
    doc.add_heading('3.2 支持的供应商', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light Grid Accent 1'
    data = [
        ('供应商', '类型', 'Base URL', '说明'),
        ('DeepSeek', '云端', 'https://api.deepseek.com/v1', '性价比高，推荐新手使用'),
        ('OpenAI', '云端', 'https://api.openai.com/v1', 'GPT-4o 系列模型'),
        ('Ollama', '本地', 'http://localhost:11434/v1', '免费，无需联网，需要安装 Ollama'),
        ('llama.cpp', '本地', 'http://127.0.0.1:8080/v1', '高性能本地推理，适合微调模型'),
        ('LM Studio', '本地', 'http://127.0.0.1:1234/v1', 'GUI 管理，适合新手'),
    ]
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_heading('3.3 配置步骤', level=2)
    doc.add_paragraph('步骤 1：打开设置')
    doc.add_paragraph('在左侧导航栏点击 "设置"（齿轮图标）', style='List Bullet')
    doc.add_paragraph('点击 "🔌 AI 供应商" 标签', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('步骤 2：添加供应商')
    doc.add_paragraph('点击 "+ 添加供应商" 按钮', style='List Bullet')
    doc.add_paragraph('选择预设模板（如 Ollama、llama.cpp 等）', style='List Bullet')
    doc.add_paragraph('填写名称、Base URL、API Key（本地模型可留空）', style='List Bullet')
    doc.add_paragraph('填写模型列表（逗号分隔，如：qwen3, llama3.2）', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('步骤 3：测试连接')
    doc.add_paragraph('点击供应商卡片上的 "测试连接" 按钮', style='List Bullet')
    doc.add_paragraph('系统会自动发现可用模型并更新列表', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('步骤 4：启用/禁用供应商')
    doc.add_paragraph('点击 "已启用/已禁用" 按钮切换状态', style='List Bullet')
    
    doc.add_heading('3.4 配置 Agent 路由', level=2)
    doc.add_paragraph(
        'Agent 路由决定了每个写作阶段使用哪个供应商和模型。'
    )
    doc.add_paragraph('步骤 1：点击 "🔀 Agent 路由" 按钮', style='List Bullet')
    doc.add_paragraph('步骤 2：为每个 Agent 选择供应商和模型', style='List Bullet')
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('Agent', '建议配置'),
        ('规划师 (Planner)', '云端供应商（DeepSeek/OpenAI），低 Temperature 0.3'),
        ('写手 (Writer)', '本地模型（llama.cpp/Ollama）或云端，高 Temperature 0.8'),
        ('深度审计 (DeepAudit)', '云端供应商（需要强推理能力），低 Temperature 0.1'),
        ('润色师 (Polisher)', '本地微调模型，中 Temperature 0.3'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph('')
    doc.add_paragraph('步骤 3：点击 "保存路由配置"')
    
    doc.add_heading('3.5 推荐配置方案', level=2)
    doc.add_paragraph('方案 A：纯云端模式（最简单）')
    p = doc.add_paragraph('配置 DeepSeek 供应商', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('所有 Agent 都使用 DeepSeek', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('月费约 $2-4', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    
    doc.add_paragraph('方案 B：混合模式（推荐）')
    p = doc.add_paragraph('云端：DeepSeek 用于规划和审计', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('本地：Ollama/llama.cpp 用于写手和润色', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('降低成本，保护隐私', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    
    doc.add_paragraph('方案 C：纯本地模式（零费用）')
    p = doc.add_paragraph('安装 Ollama 或 LM Studio', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('下载 Qwen3.6-35B 或类似模型', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('所有 Agent 都使用本地模型', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('需要 16GB+ 内存或独立显卡', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    
    doc.add_page_break()
    
    # ==================== CHAPTER 4 ====================
    doc.add_heading('第四章 创建作品与风格注入', level=1)
    
    doc.add_heading('4.1 创建新书', level=2)
    doc.add_paragraph('步骤 1：进入书架页')
    doc.add_paragraph('打开 http://localhost:3000', style='List Bullet')
    doc.add_paragraph('点击 "[+ 创建新书]" 按钮', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('步骤 2：完成 6 步问卷')
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('步骤', '填写内容'),
        ('1', '书名：给你的小说起个名字'),
        ('2', '题材：玄幻修仙 / 都市重生 / 科幻末世 / 悬疑灵异 / 古代言情'),
        ('3', '核心设定：一句话概括故事主线'),
        ('4', '主要角色：姓名、身份、性格特点'),
        ('5', '世界观：修炼体系、地理环境、势力分布'),
        ('6', '写作目标：预计字数、更新频率'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph('')
    doc.add_paragraph('步骤 3：点击 "创建"')
    
    doc.add_heading('4.2 上传风格样本（可选）', level=2)
    doc.add_paragraph('如果你有之前的网文作品，可以上传 3-5 章让 AI 学习你的写作风格：')
    doc.add_paragraph('进入作品工作台', style='List Bullet')
    doc.add_paragraph('点击左侧 "风格管理" 标签', style='List Bullet')
    doc.add_paragraph('上传你的旧稿（.txt 或 .md 格式）', style='List Bullet')
    doc.add_paragraph('点击 "提取风格指纹"', style='List Bullet')
    doc.add_paragraph('等待 1-2 分钟，系统会自动生成 style_fingerprint.json', style='List Bullet')
    
    doc.add_heading('4.3 放入参考书籍（可选）', level=2)
    doc.add_paragraph('如果你想让 AI 学习同类题材的写作技巧：')
    doc.add_paragraph('将 TXT 文件放入 data/raw-books/玄幻/ 目录', style='List Bullet')
    doc.add_paragraph('运行命令：', style='List Bullet')
    add_code_block(doc, 'pnpm run knowledge:process')
    
    doc.add_page_break()
    
    # ==================== CHAPTER 5 ====================
    doc.add_heading('第五章 日常创作流程', level=1)
    
    doc.add_heading('5.1 启动系统', level=2)
    doc.add_paragraph('运行 start.bat 或 pnpm dev', style='List Bullet')
    doc.add_paragraph('访问 http://localhost:3000', style='List Bullet')
    
    doc.add_heading('5.2 打开作品', level=2)
    doc.add_paragraph('在书架点击作品 → 进入工作台', style='List Bullet')
    
    doc.add_heading('5.3 写下一章（核心流程）', level=2)
    doc.add_paragraph('点击 [写下一章] 按钮，系统自动执行以下流水线：')
    
    doc.add_paragraph('')
    doc.add_paragraph('阶段 1：Planner 生成章纲（约 2 秒）')
    doc.add_paragraph('AI 读取卷纲 + 块纲 → 生成章纲 + 场景卡', style='List Bullet')
    doc.add_paragraph('系统弹出大纲审批面板', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('阶段 2：作者审批大纲（人工决策点 1）')
    doc.add_paragraph('审阅大纲，决定是否调整', style='List Bullet')
    doc.add_paragraph('选项：[批准] [修改大纲] [跳过本章]', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('阶段 3：Writer 流式生成正文（约 15-30 秒）')
    doc.add_paragraph('AI 装配全文记忆 + 风格指纹 → 分段生成', style='List Bullet')
    doc.add_paragraph('用户在此期间看到内容逐步出现（流式输出）', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('阶段 4：FastAudit 快速检查（约 3 秒）')
    doc.add_paragraph('12 项规则引擎检查，零 LLM 调用', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('阶段 5：DeepAudit 深度审计（约 8-15 秒）')
    doc.add_paragraph('15 维度聚焦式 LLM 审计', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('阶段 6：作者审计后介入（人工决策点 2）')
    doc.add_paragraph('查看审计报告，决定如何处理', style='List Bullet')
    doc.add_paragraph('选项：[继续] [人工修改] [标记问题] [重写某段]', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('阶段 7：Analyst 提取事实 + 更新状态（约 5 秒）')
    doc.add_paragraph('阶段 8：Polisher 去 AI 味 + 润色（约 5 秒）')
    doc.add_paragraph('阶段 9：更新全文记忆（约 1 秒）')
    
    doc.add_paragraph('')
    doc.add_paragraph('阶段 10：最终审阅（人工决策点 3）')
    doc.add_paragraph('审阅终稿，决定发布或修改', style='List Bullet')
    doc.add_paragraph('选项：[通过发布] [手动修改] [重写]', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('总耗时：约 35-60 秒')
    doc.add_paragraph('用户感知：约 15 秒（Writer 流式输出期间有内容显示）')
    
    doc.add_heading('5.4 导出发布', level=2)
    doc.add_paragraph('点击 [导出] → 选择格式（TXT/DOCX/PDF/EPUB）→ 下载', style='List Bullet')
    doc.add_paragraph('或点击 [生成封面] → 获取封面图', style='List Bullet')
    doc.add_paragraph('或点击 [导出剧本] → 获取短剧格式', style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== CHAPTER 6 ====================
    doc.add_heading('第六章 本地模型部署（进阶）', level=1)
    
    doc.add_heading('6.1 安装 Ollama', level=2)
    doc.add_paragraph('访问 https://ollama.ai 下载安装', style='List Bullet')
    doc.add_paragraph('打开命令行，下载模型：', style='List Bullet')
    add_code_block(doc, 'ollama pull qwen3')
    
    doc.add_paragraph('')
    doc.add_heading('6.2 安装 LM Studio', level=2)
    doc.add_paragraph('访问 https://lmstudio.ai 下载安装', style='List Bullet')
    doc.add_paragraph('在 GUI 中搜索并下载模型', style='List Bullet')
    doc.add_paragraph('启动本地服务器（默认端口 1234）', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_heading('6.3 安装 llama.cpp', level=2)
    doc.add_paragraph('访问 https://github.com/ggerganov/llama.cpp 下载预编译版本', style='List Bullet')
    doc.add_paragraph('运行：', style='List Bullet')
    add_code_block(doc, 'server.exe -m model.gguf -c 4096')
    
    doc.add_paragraph('')
    doc.add_heading('6.4 在 NovelForge 中配置本地模型', level=2)
    doc.add_paragraph('进入设置 → AI 供应商', style='List Bullet')
    doc.add_paragraph('添加供应商 → 选择对应预设（Ollama/llama.cpp/LM Studio）', style='List Bullet')
    doc.add_paragraph('Base URL 自动填充，无需 API Key', style='List Bullet')
    doc.add_paragraph('点击测试连接验证', style='List Bullet')
    doc.add_paragraph('在 Agent 路由中将 Writer/Polisher 指向本地供应商', style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== CHAPTER 7 ====================
    doc.add_heading('第七章 微调训练（可选）', level=1)
    
    doc.add_heading('7.1 微调是什么？', level=2)
    doc.add_paragraph(
        '微调是通过你的作品数据训练一个专属模型，让 AI 更懂你的写作风格。'
        '这需要 GPU（RTX 3070 或以上）和较长时间（24-48 小时）。'
    )
    
    doc.add_heading('7.2 微调前准备', level=2)
    doc.add_paragraph('确认 GPU 可用：', style='List Bullet')
    add_code_block(doc, 'nvidia-smi')
    
    doc.add_paragraph('')
    doc.add_paragraph('准备训练数据（如有参考书籍）：', style='List Bullet')
    add_code_block(doc, 'pnpm run fine-tune:prepare')
    
    doc.add_heading('7.3 开始训练', level=2)
    doc.add_paragraph('运行一键脚本：', style='List Bullet')
    add_code_block(doc, 'bash scripts/fine-tune.sh full')
    
    doc.add_paragraph('')
    doc.add_paragraph('脚本会自动执行以下步骤：')
    doc.add_paragraph('检查环境和 GPU', style='List Bullet')
    doc.add_paragraph('安装 Python 依赖（torch, transformers, peft 等）', style='List Bullet')
    doc.add_paragraph('生成训练数据', style='List Bullet')
    doc.add_paragraph('验证数据质量', style='List Bullet')
    doc.add_paragraph('开始训练（约 24-48 小时）', style='List Bullet')
    
    doc.add_paragraph('')
    doc.add_paragraph('训练完成后，模型保存到 models/novelforge-lora/')
    doc.add_paragraph('在 AI 供应商中配置 llama.cpp 指向该模型', style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== CHAPTER 8 ====================
    doc.add_heading('第八章 差异化功能', level=1)
    
    doc.add_heading('8.1 一键封面生成', level=2)
    doc.add_paragraph('功能：根据小说内容自动生成封面图')
    doc.add_paragraph('流程：')
    p = doc.add_paragraph('读取 MASTER_SETTING（标题、题材、核心设定）', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('V4-Flash 生成封面提示词', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('调用 SD API 生成图片', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('保存到作品目录下的 cover.png', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph('')
    doc.add_paragraph('用户操作：')
    p = doc.add_paragraph('点击 [生成封面] → 等待 10 秒 → 预览 → 确认/重新生成', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('成本：~$0.01/张', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    
    doc.add_heading('8.2 短剧剧本导出', level=2)
    doc.add_paragraph('功能：将章节转换为短剧/漫剧剧本格式')
    doc.add_paragraph('转换规则：')
    p = doc.add_paragraph('章节内容 → 场景分割', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('叙述文本 → 旁白字幕', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('对话 → 角色台词', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('动作描写 → 镜头指示', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('情绪描写 → 配乐/音效提示', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    doc.add_paragraph('')
    doc.add_paragraph('用户操作：')
    p = doc.add_paragraph('点击 [导出剧本] → 选择格式 → 下载', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    p = doc.add_paragraph('成本：$0（规则引擎）', style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.75)
    
    doc.add_page_break()
    
    # ==================== CHAPTER 9 ====================
    doc.add_heading('第九章 常见问题与故障排查', level=1)
    
    doc.add_heading('9.1 服务无法启动', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('问题', '解决方法'),
        ('端口 3000/3001 被占用', '运行 stop.bat 停止服务，或修改 .env 中的 PORT 配置'),
        ('Node.js 版本过低', '升级到 v20 或以上：nvm install 20'),
        ('pnpm install 失败', '删除 node_modules 和 pnpm-lock.yaml，重新运行 pnpm install'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_heading('9.2 API 调用失败', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('问题', '解决方法'),
        ('DEEPSEEK_API_KEY 无效', '检查 .env 文件中的 API Key 是否正确，前往 https://platform.deepseek.com 重新获取'),
        ('连接超时', '检查网络连接，或使用本地模型作为备用'),
        ('模型不存在', '在 AI 供应商中点击"测试连接"自动发现可用模型'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_heading('9.3 本地模型无法连接', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('问题', '解决方法'),
        ('Ollama 未启动', '运行 ollama serve 或在后台启动 Ollama'),
        ('llama.cpp 服务未运行', '运行 server.exe 启动服务'),
        ('LM Studio 未启动服务器', '在 LM Studio GUI 中点击 "Start Server"'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_heading('9.4 生成内容质量不佳', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('问题', '解决方法'),
        ('AI 味太重', '在设置中调整 Temperature 参数；使用本地微调模型'),
        ('风格不一致', '上传更多风格样本；调整风格指纹'),
        ('情节逻辑问题', '检查大纲质量；在审计后介入时人工修改'),
        ('角色名错误', '完善角色档案；使用 DeepAudit 检查'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_page_break()
    
    # ==================== CHAPTER 10 ====================
    doc.add_heading('第十章 附录：一键部署脚本', level=1)
    
    doc.add_heading('10.1 项目文件说明', level=2)
    doc.add_paragraph('项目根目录包含以下关键文件和脚本：', style='List Bullet')
    
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('文件/目录', '说明'),
        ('start.bat', '一键启动前后端服务（Windows）'),
        ('stop.bat', '一键停止所有服务（Windows）'),
        ('restart.bat', '重启服务（先停后启）'),
        ('.env', '环境变量配置文件（需手动创建）'),
        ('.env.example', '环境变量模板'),
        ('scripts/init.ts', '项目初始化脚本'),
        ('scripts/deploy.sh', 'Linux/macOS 部署脚本'),
        ('scripts/fine-tune.sh', '微调训练脚本'),
        ('scripts/fine-tune.py', '微调 Python 脚本'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_heading('10.2 start.bat 脚本详解', level=2)
    doc.add_paragraph('start.bat 是 Windows 下的一键启动脚本，功能：', style='List Bullet')
    doc.add_paragraph('启动后端服务（tsx + Hono，端口 3001）', style='List Bullet')
    doc.add_paragraph('启动前端开发服务器（Vite，端口 3000）', style='List Bullet')
    doc.add_paragraph('自动打开浏览器访问 http://localhost:3000', style='List Bullet')
    doc.add_paragraph('记录进程 PID 以便 stop.bat 精准停止', style='List Bullet')
    
    doc.add_heading('10.3 stop.bat 脚本详解', level=2)
    doc.add_paragraph('stop.bat 是一键停止脚本，功能：', style='List Bullet')
    doc.add_paragraph('根据 PID 文件精准停止后端和前端进程', style='List Bullet')
    doc.add_paragraph('如果 PID 文件不存在，扫描端口 3000/3001 强制终止', style='List Bullet')
    doc.add_paragraph('清理 .pids 目录', style='List Bullet')
    
    doc.add_heading('10.4 deploy.sh 脚本（Linux/macOS）', level=2)
    doc.add_paragraph('deploy.sh 是 Linux/macOS 下的部署脚本，功能包括：', style='List Bullet')
    doc.add_paragraph('检查 Node.js 和 npm 是否安装', style='List Bullet')
    doc.add_paragraph('安装前后端依赖', style='List Bullet')
    doc.add_paragraph('创建必要目录（workspace, data 等）', style='List Bullet')
    doc.add_paragraph('复制 .env.example 为 .env', style='List Bullet')
    
    doc.add_heading('10.5 fine-tune.sh 脚本（微调训练）', level=2)
    doc.add_paragraph('fine-tune.sh 支持以下模式：', style='List Bullet')
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    data = [
        ('模式', '说明'),
        ('check', '检查环境和 GPU'),
        ('generate', '生成和验证训练数据'),
        ('train', '验证数据并开始训练'),
        ('dpo', '运行 DPO 偏好学习训练'),
        ('full', '完整流程（检查 + 生成 + 训练）'),
        ('status', '查看现有训练状态'),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i].cells[0].text = col1
        table.rows[i].cells[1].text = col2
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph('')
    doc.add_paragraph('使用方法：', style='List Bullet')
    add_code_block(doc, 'bash scripts/fine-tune.sh full')
    
    doc.add_heading('10.6 项目目录结构', level=2)
    add_code_block(doc, 'novelforge/')
    add_code_block(doc, '├── src/              # 后端源码')
    add_code_block(doc, '│   ├── agents/       # AI 智能体')
    add_code_block(doc, '│   ├── core/         # 核心管线')
    add_code_block(doc, '│   ├── state/        # 状态管理')
    add_code_block(doc, '│   └── api/          # HTTP API')
    add_code_block(doc, '├── studio/           # 前端应用')
    add_code_block(doc, '│   └── src/')
    add_code_block(doc, '│       └── pages/    # 页面组件')
    add_code_block(doc, '├── workspace/        # 工作区数据')
    add_code_block(doc, '├── data/             # 数据目录')
    add_code_block(doc, '│   ├── raw-books/    # 参考书籍')
    add_code_block(doc, '│   └── training/     # 训练数据')
    add_code_block(doc, '├── scripts/          # 工具脚本')
    add_code_block(doc, '│   ├── init.ts       # 初始化')
    add_code_block(doc, '│   ├── deploy.sh     # 部署')
    add_code_block(doc, '│   └── fine-tune.sh  # 微调')
    add_code_block(doc, '├── templates/        # 题材模板')
    add_code_block(doc, '├── start.bat         # 启动脚本')
    add_code_block(doc, '├── stop.bat          # 停止脚本')
    add_code_block(doc, '├── .env              # 环境变量')
    add_code_block(doc, '└── .env.example      # 环境变量模板')
    
    doc.add_page_break()
    
    # ==================== SUMMARY ====================
    doc.add_heading('总结', level=1)
    doc.add_paragraph(
        'NovelForge 是一个强大的 AI 辅助创作工具，通过以下步骤即可开始使用：\n\n'
        '1. 安装 Node.js 和 pnpm（约 15 分钟）\n'
        '2. 克隆项目并安装依赖（约 10 分钟）\n'
        '3. 配置 API Key（约 5 分钟）\n'
        '4. 启动服务并创建作品（约 10 分钟）\n'
        '5. 开始日常创作！（每章约 35-60 秒）\n\n'
        '核心优势：\n'
        '• 全文记忆：50 章+一致性\n'
        '• 风格迁移：学习你的写作风格\n'
        '• 多供应商：云端+本地灵活搭配\n'
        '• 一键部署：start.bat 即可运行\n'
        '• 月度成本低：$2-4/月\n\n'
        '记住：AI 执行，人工决策。你始终掌握创作的控制权。'
    )
    
    # Save document
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'NovelForge_用户操作手册_v3.5.docx')
    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    print(f'文档已生成：{output_path}')
    print(f'文件大小：{file_size / 1024 / 1024:.2f} MB')


if __name__ == '__main__':
    main()
